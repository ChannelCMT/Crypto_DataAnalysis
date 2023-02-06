import os
from docx import Document
from docx.shared import Inches
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE

os.system("python ./volumeProfileAuto.py 2023-01-01 2023-01-31")

document = Document()
styles = document.styles

document.add_heading('Monthly Report of DcBot', 0)

def creatTableFromDataFrame(df, table):
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]
        if j == 0:
            table.cell(0,j).width = Inches(3)
        else:
            table.cell(0,j).width = Inches(1)
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])

pVolatility = document.add_paragraph()
pVolatility.add_run('1. 过去一个月每周的成交量密集成交区变化。').bold = True
document.add_picture('btcVolumeProfile.png', width=Inches(6.2), height=Inches(2))
document.add_paragraph('-上图是BTC本月每周的密集成交区分布图')
document.add_picture('ethVolumeProfile.png', width=Inches(6.2), height=Inches(2))
document.add_paragraph('-上图是ETH本月每周的密集成交区分布图')

data1Excel = pd.read_excel('btcVolumeProfileDf.xlsx')
data1Df = data1Excel.rename(columns={'Unnamed: 0': 'POC_SAR'})
document.add_paragraph('表格显示BTC过去一个月中每周的密集成交区间')
data1Table = document.add_table(data1Df.shape[0]+1,data1Df.shape[1],style = 'Light Grid Accent 1')
creatTableFromDataFrame(data1Df, data1Table)

data2Excel = pd.read_excel('ethVolumeProfileDf.xlsx')
data2Df = data2Excel.rename(columns={'Unnamed: 0': 'POC_SAR'})
document.add_paragraph('表格显示ETH过去一个月中每周的密集成交区间')
data2Table = document.add_table(data2Df.shape[0]+1,data2Df.shape[1],style = 'Light Grid Accent 1')
creatTableFromDataFrame(data2Df, data2Table)

document.add_page_break()

document.save('volumeProfileMonthly.docx')