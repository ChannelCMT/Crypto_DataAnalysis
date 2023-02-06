import os
from docx import Document
from docx.shared import Inches
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE

os.system("python ./priceAutoCode.py 2023-01-27 2023-02-03")

document = Document()
styles = document.styles

document.add_heading('Monthly Report of DcBot', 0)

def creatTableFromDataFrame(df, table):
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]
        if j == 0:
            table.cell(0,j).width = Inches(3)
        else:
            table.cell(0,j).width = Inches(1)
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])

pVolatility = document.add_paragraph()
pVolatility.add_run('1. 本月的价格走势与每周的变化率。').bold = True
document.add_picture('coin1_price.png', width=Inches(6.2), height=Inches(2))
document.add_paragraph('-上图是BTC本月的价格图')
document.add_picture('coin2_price.png', width=Inches(6.2), height=Inches(2))
document.add_paragraph('-上图是ETH本月的价格图')

dataExcel = pd.read_excel('PricePctChange.xlsx')
dataDf = dataExcel.rename(columns={'Unnamed: 0': 'pctChange'})
document.add_paragraph('表格显示过去一个月中每周的价格变化率')
dataTable = document.add_table(dataDf.shape[0]+1,dataDf.shape[1],style = 'Light Grid Accent 1')
creatTableFromDataFrame(dataDf, dataTable)

document.add_page_break()

document.save('priceChange.docx')